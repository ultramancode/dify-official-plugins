"""Abstract interface for document loader implementations."""

import logging
import mimetypes
import re
import uuid
from io import BytesIO
from urllib.parse import urlparse
from xml.etree import ElementTree

import requests
from dify_plugin import Tool
from docx import Document as DocxDocument

from tools.extractor_base import BaseExtractor
from tools.document import Document, ExtractorResult

logger = logging.getLogger(__name__)


class WordExtractor(BaseExtractor):
    """Load docx files.

    Args:
        tool: Tool instance
        file_bytes: file bytes
        file_name: file name
        encoding: encoding
        autodetect_encoding: autodetect encoding
    """

    def __init__(self, tool: Tool, file_bytes: bytes, file_name: str):
        """Initialize with file path."""
        self._file_bytes = file_bytes
        self._file_name = file_name
        self._tool = tool

    def extract(self) -> ExtractorResult:
        """Load given path as single page."""
        content, img_list = self.parse_docx(self._file_bytes)
        return ExtractorResult(
            md_content=content,
            documents=[
                Document(page_content=content, metadata={"source": self._file_name})
            ],
            img_list=img_list,
            origin_result=None
        )

    @staticmethod
    def _is_valid_url(url: str) -> bool:
        """Check if the url is valid."""
        parsed = urlparse(url)
        return bool(parsed.netloc) and bool(parsed.scheme)

    def _extract_images_from_docx(self, doc):
        image_count = 0
        image_map = {}
        img_list = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                if rel.is_external:
                    url = rel.target_ref
                    if not self._is_valid_url(url):
                        continue
                    response = requests.get(url)
                    if response.status_code == 200:
                        image_ext = mimetypes.guess_extension(
                            response.headers["Content-Type"]
                        )
                        if image_ext is None:
                            continue
                        file_uuid = str(uuid.uuid4())
                        file_name = file_uuid + "." + image_ext
                        mime_type, _ = mimetypes.guess_type(file_name)

                        file_res = self._tool.session.file.upload(
                            file_name, response.content, mime_type
                        )
                    else:
                        continue
                else:
                    image_ext = rel.target_ref.split(".")[-1]
                    if image_ext is None:
                        continue
                    # user uuid as file name
                    file_uuid = str(uuid.uuid4())
                    file_name = file_uuid + "." + image_ext
                    mime_type, _ = mimetypes.guess_type(file_name)

                    file_res = self._tool.session.file.upload(
                        file_name, rel.target_part.blob, mime_type
                    )

                image_map[rel.target_part] = f"![image]({file_res.preview_url})"
                img_list.append(file_res)

        return image_map, img_list

    def _table_to_markdown(self, table, image_map):
        markdown = ""
        # calculate the total number of columns
        total_cols = max(len(row.cells) for row in table.rows)

        header_row = table.rows[0]
        headers = self._parse_row(header_row, image_map, total_cols)
        markdown += "| " + " | ".join(headers) + " |\n"
        markdown += "| " + " | ".join(["---"] * total_cols) + " |\n"

        for row in table.rows[1:]:
            row_cells = self._parse_row(row, image_map, total_cols)
            markdown += "| " + " | ".join(row_cells) + " |\n"

        return markdown

    def _parse_row(self, row, image_map, total_cols):
        # Initialize a row, all of which are empty by default
        row_cells = [""] * total_cols
        col_index = 0
        for cell in row.cells:
            # make sure the col_index is not out of range
            while col_index < total_cols and row_cells[col_index] != "":
                col_index += 1
            # if col_index is out of range the loop is jumped
            if col_index >= total_cols:
                break
            cell_content = self._parse_cell(cell, image_map).strip()
            cell_colspan = cell.grid_span or 1
            for i in range(cell_colspan):
                if col_index + i < total_cols:
                    row_cells[col_index + i] = cell_content if i == 0 else ""
            col_index += cell_colspan
        return row_cells

    def _parse_cell(self, cell, image_map):
        cell_content = []
        for paragraph in cell.paragraphs:
            parsed_paragraph = self._parse_cell_paragraph(paragraph, image_map)
            if parsed_paragraph:
                cell_content.append(parsed_paragraph)
        unique_content = list(dict.fromkeys(cell_content))
        return " ".join(unique_content)

    def _parse_cell_paragraph(self, paragraph, image_map):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    image_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if not image_id:
                        continue
                    image_part = paragraph.part.rels[image_id].target_part

                    if image_part in image_map:
                        image_link = image_map[image_part]
                        paragraph_content.append(image_link)
            else:
                paragraph_content.append(run.text)
        return "".join(paragraph_content).strip()

    def _parse_paragraph(self, paragraph, image_map):
        paragraph_content = []
        for run in paragraph.runs:
            if run.element.xpath(".//a:blip"):
                for blip in run.element.xpath(".//a:blip"):
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if embed_id:
                        rel_target = run.part.rels[embed_id].target_ref
                        if rel_target in image_map:
                            paragraph_content.append(image_map[rel_target])
            if run.text.strip():
                paragraph_content.append(run.text.strip())
        return " ".join(paragraph_content) if paragraph_content else ""

    def parse_docx(self, file_bytes):
        doc = DocxDocument(BytesIO(file_bytes))

        content = []

        image_map, img_list = self._extract_images_from_docx(doc)

        hyperlinks_url = None
        url_pattern = re.compile(r"http://[^\s+]+//|https://[^\s+]+")
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text and hyperlinks_url:
                    result = f"  [{run.text}]({hyperlinks_url})  "
                    run.text = result
                    hyperlinks_url = None
                if "HYPERLINK" in run.element.xml:
                    try:
                        xml = ElementTree.XML(run.element.xml)
                        x_child = [c for c in xml.iter() if c is not None]
                        for x in x_child:
                            if x_child is None:
                                continue
                            if x.tag.endswith("instrText"):
                                if x.text is None:
                                    continue
                                for i in url_pattern.findall(x.text):
                                    hyperlinks_url = str(i)
                    except Exception:
                        logger.exception("Failed to parse HYPERLINK xml")

        def parse_paragraph(paragraph):
            paragraph_content = []
            for run in paragraph.runs:
                if (
                    hasattr(run.element, "tag")
                    and isinstance(run.element.tag, str)
                    and run.element.tag.endswith("r")
                ):
                    drawing_elements = run.element.findall(
                        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                    )
                    for drawing in drawing_elements:
                        blip_elements = drawing.findall(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                        )
                        for blip in blip_elements:
                            embed_id = blip.get(
                                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                            )
                            if embed_id:
                                image_part = doc.part.related_parts.get(embed_id)
                                if image_part in image_map:
                                    paragraph_content.append(image_map[image_part])
                if run.text.strip():
                    paragraph_content.append(run.text.strip())
            return "".join(paragraph_content) if paragraph_content else ""

        paragraphs = doc.paragraphs.copy()
        tables = doc.tables.copy()
        for element in doc.element.body:
            if hasattr(element, "tag"):
                if isinstance(element.tag, str) and element.tag.endswith(
                    "p"
                ):  # paragraph
                    para = paragraphs.pop(0)
                    parsed_paragraph = parse_paragraph(para)
                    if parsed_paragraph.strip():
                        content.append(parsed_paragraph)
                    else:
                        content.append("\n")
                elif isinstance(element.tag, str) and element.tag.endswith(
                    "tbl"
                ):  # table
                    table = tables.pop(0)
                    content.append(self._table_to_markdown(table, image_map))
        return "\n".join(content), img_list
